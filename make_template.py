"""
Script to create the HBL Sales Enablement DOCX template.
This version matches the actual Pydantic models provided.
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "templates")


def create_hbl_sales_enablement_template():
    """Create the HBL Sales Enablement template matching Pydantic models."""
    
    if not os.path.exists(TEMPLATE_DIR):
        os.makedirs(TEMPLATE_DIR)
    
    doc = Document()
    
    # ==========================================
    # HEADER
    # ==========================================
    title = doc.add_heading("Hypothekarbank Lenzburg", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading("Kundenanalyse - Sales Enablement Report", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Erstellungsdatum: ").bold = True
    p.add_run("{{ report_date }}")
    p.add_run("    |    ")
    p.add_run("Steuerjahr: ").bold = True
    p.add_run("{{ steuererklaerung_data.tax_year }}")
    
    # ==========================================
    # PART 1: STRATEGIC ANALYSIS & CLIENT RATING
    # ==========================================
    doc.add_heading("Teil 1: Strategische Analyse & Kundenrating", level=1)
    
    # Rating Box
    doc.add_heading("Client Opportunity Score", level=2)
    
    # Main rating display
    rating_table = doc.add_table(rows=2, cols=4)
    rating_table.style = "Table Grid"
    
    # Header row
    headers = ["Kundenrating", "Gesamtscore", "Potenzial-Score", "Trigger-Score"]
    for i, header in enumerate(headers):
        rating_table.rows[0].cells[i].text = header
        rating_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    # Value row
    rating_table.rows[1].cells[0].text = "{{ rating_result.rating.value }} - {{ rating_result.rating_label }}"
    rating_table.rows[1].cells[1].text = "{{ rating_result.score.total_score|round(1) }}/100"
    rating_table.rows[1].cells[2].text = "{{ rating_result.score.potential_score|round(1) }}/100"
    rating_table.rows[1].cells[3].text = "{{ rating_result.score.trigger_score|round(1) }}/100"
    
    # Rating Rationale
    doc.add_paragraph()
    doc.add_heading("Bewertungsbegründung", level=3)
    doc.add_paragraph("{{ rating_result.rating_rationale }}")
    
    # Recommended Action
    p = doc.add_paragraph()
    p.add_run("Empfohlene Massnahme: ").bold = True
    p.add_run("{{ rating_result.recommended_action }}")
    
    # Focus Areas
    p = doc.add_paragraph()
    p.add_run("Fokusgebiete: ").bold = True
    doc.add_paragraph("{% for area in rating_result.focus_areas %}• {{ area }}\n{% endfor %}")
    
    # Primary Triggers
    doc.add_heading("Identifizierte Trigger", level=3)
    doc.add_paragraph("{% for trigger in rating_result.primary_triggers %}• {{ trigger }}\n{% endfor %}")
    
    # ==========================================
    # SCORING MATRIX - Potential Dimension
    # ==========================================
    doc.add_heading("Scoring Matrix: Potenzial-Dimension", level=2)
    
    potential_table = doc.add_table(rows=1, cols=6)
    potential_table.style = "Table Grid"
    
    # Headers
    pot_headers = ["Kategorie", "Wert", "Punkte", "Max", "Gewicht", "Begründung"]
    for i, h in enumerate(pot_headers):
        potential_table.rows[0].cells[i].text = h
        potential_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    # Add loop for dynamic rows
    doc.add_paragraph("{% for comp in rating_result.score.potential_components %}")
    
    # Data row table (will be repeated)
    comp_table_pot = doc.add_table(rows=1, cols=6)
    comp_table_pot.style = "Table Grid"
    comp_table_pot.rows[0].cells[0].text = "{{ comp.category }}"
    comp_table_pot.rows[0].cells[1].text = "{{ comp.raw_value if comp.raw_value else '-' }}"
    comp_table_pot.rows[0].cells[2].text = "{{ comp.points }}"
    comp_table_pot.rows[0].cells[3].text = "{{ comp.max_points }}"
    comp_table_pot.rows[0].cells[4].text = "{{ comp.weight_percent }}%"
    comp_table_pot.rows[0].cells[5].text = "{{ comp.reasoning }}"
    
    doc.add_paragraph("{% endfor %}")
    
    # ==========================================
    # SCORING MATRIX - Trigger Dimension
    # ==========================================
    doc.add_heading("Scoring Matrix: Trigger-Dimension", level=2)
    
    trigger_table = doc.add_table(rows=1, cols=6)
    trigger_table.style = "Table Grid"
    
    # Headers
    for i, h in enumerate(pot_headers):
        trigger_table.rows[0].cells[i].text = h
        trigger_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    # Add loop for dynamic rows
    doc.add_paragraph("{% for comp in rating_result.score.trigger_components %}")
    
    # Data row table (will be repeated)
    comp_table_trig = doc.add_table(rows=1, cols=6)
    comp_table_trig.style = "Table Grid"
    comp_table_trig.rows[0].cells[0].text = "{{ comp.category }}"
    comp_table_trig.rows[0].cells[1].text = "{{ comp.raw_value if comp.raw_value else '-' }}"
    comp_table_trig.rows[0].cells[2].text = "{{ comp.points }}"
    comp_table_trig.rows[0].cells[3].text = "{{ comp.max_points }}"
    comp_table_trig.rows[0].cells[4].text = "{{ comp.weight_percent }}%"
    comp_table_trig.rows[0].cells[5].text = "{{ comp.reasoning }}"
    
    doc.add_paragraph("{% endfor %}")
    
    # ==========================================
    # PART 2: PERSONAL MASTER DATA
    # ==========================================
    doc.add_page_break()
    doc.add_heading("Teil 2: Persönliche Stammdaten", level=1)
    
    master_table = doc.add_table(rows=6, cols=3)
    master_table.style = "Table Grid"
    
    # Headers
    master_table.rows[0].cells[0].text = "Feld"
    master_table.rows[0].cells[1].text = "Person 1"
    master_table.rows[0].cells[2].text = "Person 2"
    for cell in master_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
    
    # Data rows
    master_table.rows[1].cells[0].text = "Name"
    master_table.rows[1].cells[1].text = "{{ steuererklaerung_data.person_1.master_data.name if steuererklaerung_data.person_1.master_data else '-' }}"
    master_table.rows[1].cells[2].text = "{{ steuererklaerung_data.person_2.master_data.name if steuererklaerung_data.person_2 and steuererklaerung_data.person_2.master_data else '-' }}"
    
    master_table.rows[2].cells[0].text = "Alter"
    master_table.rows[2].cells[1].text = "{{ steuererklaerung_data.person_1.master_data.alter if steuererklaerung_data.person_1.master_data else '-' }}"
    master_table.rows[2].cells[2].text = "{{ steuererklaerung_data.person_2.master_data.alter if steuererklaerung_data.person_2 and steuererklaerung_data.person_2.master_data else '-' }}"
    
    master_table.rows[3].cells[0].text = "Geburtsdatum"
    master_table.rows[3].cells[1].text = "{{ steuererklaerung_data.person_1.master_data.geburtsdatum if steuererklaerung_data.person_1.master_data else '-' }}"
    master_table.rows[3].cells[2].text = "{{ steuererklaerung_data.person_2.master_data.geburtsdatum if steuererklaerung_data.person_2 and steuererklaerung_data.person_2.master_data else '-' }}"
    
    master_table.rows[4].cells[0].text = "Erwerbsstatus"
    master_table.rows[4].cells[1].text = "{{ steuererklaerung_data.person_1.master_data.employment_status.value if steuererklaerung_data.person_1.master_data else '-' }}"
    master_table.rows[4].cells[2].text = "{{ steuererklaerung_data.person_2.master_data.employment_status.value if steuererklaerung_data.person_2 and steuererklaerung_data.person_2.master_data else '-' }}"
    
    master_table.rows[5].cells[0].text = "Zivilstand"
    master_table.rows[5].cells[1].text = "{{ 'verheiratet' if steuererklaerung_data.verheiratet else 'ledig' }}"
    master_table.rows[5].cells[2].text = "-"
    
    # Housing situation
    p = doc.add_paragraph()
    p.add_run("Wohnsituation: ").bold = True
    p.add_run("{{ steuererklaerung_data.housing_situation.value }}")
    
    # ==========================================
    # PART 3: FINANCIAL DATA
    # ==========================================
    doc.add_heading("Teil 3: Finanzielle Datenextraktion", level=1)
    
    # 3.1 Income Analysis
    doc.add_heading("3.1 Einkommensanalyse", level=2)
    
    income_table = doc.add_table(rows=4, cols=3)
    income_table.style = "Table Grid"
    
    income_table.rows[0].cells[0].text = "Einkommensart"
    income_table.rows[0].cells[1].text = "Person 1"
    income_table.rows[0].cells[2].text = "Person 2"
    for cell in income_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
    
    income_table.rows[1].cells[0].text = "Bruttoeinkommen"
    income_table.rows[1].cells[1].text = "{{ steuererklaerung_data.person_1.haupterwerb if steuererklaerung_data.person_1.haupterwerb else '-' }}"
    income_table.rows[1].cells[2].text = "{{ steuererklaerung_data.person_2.haupterwerb if steuererklaerung_data.person_2 and steuererklaerung_data.person_2.haupterwerb else '-' }}"
    
    income_table.rows[2].cells[0].text = "Nettoeinkommen"
    income_table.rows[2].cells[1].text = "{{ steuererklaerung_data.person_1.nettoeinkommen if steuererklaerung_data.person_1.nettoeinkommen else '-' }}"
    income_table.rows[2].cells[2].text = "{{ steuererklaerung_data.person_2.nettoeinkommen if steuererklaerung_data.person_2 and steuererklaerung_data.person_2.nettoeinkommen else '-' }}"
    
    income_table.rows[3].cells[0].text = "Steuerbares Einkommen"
    income_table.rows[3].cells[1].text = "{{ steuererklaerung_data.person_1.steuerbares_einkommen if steuererklaerung_data.person_1.steuerbares_einkommen else '-' }}"
    income_table.rows[3].cells[2].text = "{{ steuererklaerung_data.person_2.steuerbares_einkommen if steuererklaerung_data.person_2 and steuererklaerung_data.person_2.steuerbares_einkommen else '-' }}"
    
    # 3.2 Asset Balance
    doc.add_heading("3.2 Vermögensbilanz", level=2)
    
    # Assets
    doc.add_paragraph().add_run("Vermögenswerte").bold = True
    
    asset_table = doc.add_table(rows=5, cols=3)
    asset_table.style = "Table Grid"
    
    asset_table.rows[0].cells[0].text = "Vermögensart"
    asset_table.rows[0].cells[1].text = "Person 1"
    asset_table.rows[0].cells[2].text = "Person 2"
    for cell in asset_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
    
    asset_table.rows[1].cells[0].text = "Bankguthaben"
    asset_table.rows[1].cells[1].text = "{{ steuererklaerung_data.person_1.bankguthaben if steuererklaerung_data.person_1.bankguthaben else '-' }}"
    asset_table.rows[1].cells[2].text = "{{ steuererklaerung_data.person_2.bankguthaben if steuererklaerung_data.person_2 and steuererklaerung_data.person_2.bankguthaben else '-' }}"
    
    asset_table.rows[2].cells[0].text = "Wertschriften"
    asset_table.rows[2].cells[1].text = "{{ steuererklaerung_data.person_1.wertschriften if steuererklaerung_data.person_1.wertschriften else '-' }}"
    asset_table.rows[2].cells[2].text = "{{ steuererklaerung_data.person_2.wertschriften if steuererklaerung_data.person_2 and steuererklaerung_data.person_2.wertschriften else '-' }}"
    
    asset_table.rows[3].cells[0].text = "Liegenschaften"
    asset_table.rows[3].cells[1].text = "{{ steuererklaerung_data.person_1.liegenschaften if steuererklaerung_data.person_1.liegenschaften else '-' }}"
    asset_table.rows[3].cells[2].text = "{{ steuererklaerung_data.person_2.liegenschaften if steuererklaerung_data.person_2 and steuererklaerung_data.person_2.liegenschaften else '-' }}"
    
    asset_table.rows[4].cells[0].text = "Versicherungen"
    asset_table.rows[4].cells[1].text = "{{ steuererklaerung_data.person_1.lebens_und_versicherungspolicen if steuererklaerung_data.person_1.lebens_und_versicherungspolicen else '-' }}"
    asset_table.rows[4].cells[2].text = "{{ steuererklaerung_data.person_2.lebens_und_versicherungspolicen if steuererklaerung_data.person_2 and steuererklaerung_data.person_2.lebens_und_versicherungspolicen else '-' }}"
    
    # Liabilities
    doc.add_paragraph()
    doc.add_paragraph().add_run("Verbindlichkeiten").bold = True
    
    liab_table = doc.add_table(rows=3, cols=3)
    liab_table.style = "Table Grid"
    
    liab_table.rows[0].cells[0].text = "Verbindlichkeitsart"
    liab_table.rows[0].cells[1].text = "Person 1"
    liab_table.rows[0].cells[2].text = "Person 2"
    for cell in liab_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
    
    liab_table.rows[1].cells[0].text = "Hypotheken"
    liab_table.rows[1].cells[1].text = "{{ steuererklaerung_data.person_1.hypothekarschulden if steuererklaerung_data.person_1.hypothekarschulden else '-' }}"
    liab_table.rows[1].cells[2].text = "{{ steuererklaerung_data.person_2.hypothekarschulden if steuererklaerung_data.person_2 and steuererklaerung_data.person_2.hypothekarschulden else '-' }}"
    
    liab_table.rows[2].cells[0].text = "Übrige Schulden"
    liab_table.rows[2].cells[1].text = "{{ steuererklaerung_data.person_1.schulden if steuererklaerung_data.person_1.schulden else '-' }}"
    liab_table.rows[2].cells[2].text = "{{ steuererklaerung_data.person_2.schulden if steuererklaerung_data.person_2 and steuererklaerung_data.person_2.schulden else '-' }}"
    
    # 3.3 Pension Indicators
    doc.add_heading("3.3 Vorsorge-Indikatoren", level=2)
    
    pension_table = doc.add_table(rows=6, cols=2)
    pension_table.style = "Table Grid"
    
    pension_table.rows[0].cells[0].text = "Säule 3a Person 1"
    pension_table.rows[0].cells[1].text = "{{ steuererklaerung_data.person_1.saeule_3a_einzahlung if steuererklaerung_data.person_1.saeule_3a_einzahlung else '-' }}"
    
    pension_table.rows[1].cells[0].text = "Säule 3a Person 2"
    pension_table.rows[1].cells[1].text = "{{ steuererklaerung_data.person_2.saeule_3a_einzahlung if steuererklaerung_data.person_2 and steuererklaerung_data.person_2.saeule_3a_einzahlung else '-' }}"
    
    pension_table.rows[2].cells[0].text = "3a voll ausgeschöpft?"
    pension_table.rows[2].cells[1].text = "{{ 'Ja' if pension_indicators.saeule_3a_voll_ausgeschoepft else 'Nein' if pension_indicators.saeule_3a_voll_ausgeschoepft is not none else '-' }}"
    
    pension_table.rows[3].cells[0].text = "PK-Einkauf erkennbar?"
    pension_table.rows[3].cells[1].text = "{{ 'Ja' if pension_indicators.pk_einkauf_erkennbar else 'Nein' }}"
    
    pension_table.rows[4].cells[0].text = "Einkommens-/Sparquoten-Indikator"
    pension_table.rows[4].cells[1].text = "{{ 'Hoch' if pension_indicators.high_income_low_savings_indicator else 'Normal' }}"
    
    pension_table.rows[5].cells[0].text = "Geschätzte Vorsorgelücke"
    pension_table.rows[5].cells[1].text = "{{ pension_indicators.estimated_pension_gap if pension_indicators.estimated_pension_gap else '-' }}"
    
    # ==========================================
    # PART 4: BUSINESS OPPORTUNITIES
    # ==========================================
    doc.add_page_break()
    doc.add_heading("Teil 4: Geschäftsmöglichkeiten für HBL", level=1)
    
    doc.add_paragraph("{% if business_opportunities %}")
    doc.add_paragraph("{% for opp in business_opportunities %}")
    
    doc.add_heading("{{ opp.urgency|upper }} - {{ opp.title }}", level=2)
    
    p = doc.add_paragraph()
    p.add_run("Typ: ").bold = True
    p.add_run("{{ opp.opportunity_type }}")
    p.add_run("    |    ")
    p.add_run("Potenzial: ").bold = True
    p.add_run("{{ opp.estimated_potential if opp.estimated_potential else '-' }}")
    
    doc.add_paragraph("{{ opp.description }}")
    
    doc.add_paragraph().add_run("Nächste Schritte:").bold = True
    doc.add_paragraph("{% for step in opp.next_steps %}• {{ step }}\n{% endfor %}")
    
    doc.add_paragraph("{% endfor %}")
    doc.add_paragraph("{% else %}")
    doc.add_paragraph("Keine spezifischen Geschäftsmöglichkeiten identifiziert.")
    doc.add_paragraph("{% endif %}")
    
    # ==========================================
    # DETECTED ASSETS
    # ==========================================
    doc.add_heading("Erkannte Vermögenswerte", level=2)
    
    doc.add_paragraph("{% if detected_assets %}")
    doc.add_paragraph("{% for asset in detected_assets %}")
    
    detected_asset_table = doc.add_table(rows=1, cols=2)
    detected_asset_table.style = "Table Grid"
    detected_asset_table.rows[0].cells[0].text = "{{ asset.asset_type }}"
    detected_asset_table.rows[0].cells[1].text = "{{ asset.estimated_value if asset.estimated_value else '-' }}"
    
    doc.add_paragraph("{% endfor %}")
    doc.add_paragraph("{% else %}")
    doc.add_paragraph("Keine zusätzlichen Vermögenswerte erkannt.")
    doc.add_paragraph("{% endif %}")
    
    # ==========================================
    # NARRATIVE SECTIONS
    # ==========================================
    doc.add_page_break()
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph("{{ summary }}")
    
    doc.add_heading("Strategische Empfehlungen für Berater", level=1)
    doc.add_paragraph("{{ strategic_recommendations }}")
    
    # ==========================================
    # FOOTER
    # ==========================================
    doc.add_paragraph()
    doc.add_paragraph("_" * 60)
    p = doc.add_paragraph()
    p.add_run("Dieses Dokument wurde automatisch generiert und dient als Verkaufsunterstützung für Kundenberater der Hypothekarbank Lenzburg.").italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.add_run("Vertraulich - Nur für internen Gebrauch").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save template
    template_path = os.path.join(TEMPLATE_DIR, "hbl_sales_enablement_template.docx")
    doc.save(template_path)
    print(f"✅ Template created: {template_path}")
    print(f"\n📝 Note: This template uses separate tables per loop iteration.")
    print(f"   The template variables now match the Pydantic models exactly.\n")
    
    return template_path


if __name__ == "__main__":
    create_hbl_sales_enablement_template()